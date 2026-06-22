from __future__ import annotations

import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
PUBLIC_DOCS = ROOT / "docs" / "public"
BUILD_DIR = ROOT / ".build" / "public_docs"

GUIDE_MD = PUBLIC_DOCS / "TOAN_AAS_HUONG_DAN_SU_DUNG_CHO_KHACH_V2.md"
TERMS_MD = PUBLIC_DOCS / "TOAN_AAS_DIEU_KHOAN_CHINH_SACH_DICH_VU_V2.md"
GUIDE_DOCX = ROOT / "TOAN_AAS_HUONG_DAN_SU_DUNG_CHO_KHACH_V2.docx"
TERMS_DOCX = BUILD_DIR / "TOAN_AAS_DIEU_KHOAN_CHINH_SACH_DICH_VU_V2.docx"
GUIDE_PREVIEW_PDF = BUILD_DIR / "TOAN_AAS_HUONG_DAN_SU_DUNG_CHO_KHACH_V2_preview.pdf"
TERMS_PDF = ROOT / "TOAN_AAS_DIEU_KHOAN_CHINH_SACH_DICH_VU_V2.pdf"

GREEN = RGBColor(0x08, 0x73, 0x4F)
DARK_GREEN = RGBColor(0x05, 0x3B, 0x2B)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x18, 0x24, 0x20)
MUTED = RGBColor(0x5B, 0x68, 0x62)
LIGHT_GREEN = "E7F4EE"
LIGHT_BLUE = "E8EEF5"


def set_run_font(run, *, size=11, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_page(doc, running_label):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(running_label), size=8.5, color=MUTED, bold=True)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.paragraph_format.space_before = Pt(0)
        footer_p.paragraph_format.space_after = Pt(0)
        add_page_number(footer_p)


def configure_styles(doc, preset):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "guide" else 1.1

    heading_specs = {
        "Heading 1": (16, GREEN if preset == "guide" else BLUE, 18 if preset == "guide" else 16, 10 if preset == "guide" else 8),
        "Heading 2": (13, GREEN if preset == "guide" else BLUE, 14 if preset == "guide" else 12, 7 if preset == "guide" else 6),
        "Heading 3": (12, DARK_GREEN if preset == "guide" else DARK_BLUE, 10 if preset == "guide" else 8, 5 if preset == "guide" else 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "guide" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "guide" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "guide" else 8)
        style.paragraph_format.line_spacing = 1.25 if preset == "guide" else 1.167


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline_runs(paragraph, text):
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), name="Consolas", size=10, color=DARK_GREEN)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            set_run_font(paragraph.add_run(f"{label} ({url})"), color=GREEN, bold=True)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]))


def add_title_block(doc, *, title, subtitle, metadata, preset):
    if preset == "guide":
        kicker = doc.add_paragraph()
        kicker.paragraph_format.space_before = Pt(12)
        kicker.paragraph_format.space_after = Pt(2)
        set_run_font(kicker.add_run("TÀI LIỆU KHÁCH HÀNG"), size=10, color=GREEN, bold=True)
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(6)
        set_run_font(title_p.add_run(title), size=28, color=DARK_GREEN, bold=True)
        subtitle_p = doc.add_paragraph()
        subtitle_p.paragraph_format.space_after = Pt(18)
        set_run_font(subtitle_p.add_run(subtitle), size=13, color=MUTED)
        table = doc.add_table(rows=len(metadata), cols=2)
        set_table_geometry(table, [2700, 6660])
        for row, (label, value) in zip(table.rows, metadata):
            shade_cell(row.cells[0], LIGHT_GREEN)
            set_run_font(row.cells[0].paragraphs[0].add_run(label), size=10, color=DARK_GREEN, bold=True)
            set_run_font(row.cells[1].paragraphs[0].add_run(value), size=10, color=INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    else:
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(12)
        title_p.paragraph_format.space_after = Pt(4)
        set_run_font(title_p.add_run(title.upper()), size=23, color=INK, bold=True)
        subtitle_p = doc.add_paragraph()
        subtitle_p.paragraph_format.space_after = Pt(16)
        set_run_font(subtitle_p.add_run(subtitle), size=13, color=MUTED)
        for label, value in metadata:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(f"{label}: "), bold=True)
            set_run_font(p.add_run(value))
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(12)
        note.paragraph_format.space_after = Pt(12)
        note.paragraph_format.left_indent = Inches(0.15)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), LIGHT_BLUE)
        note._p.get_or_add_pPr().append(shade)
        set_run_font(note.add_run("Tài liệu mô tả chính sách vận hành, không thay thế tư vấn pháp lý, kế toán hoặc thuế chuyên nghiệp."), size=10.5, color=DARK_BLUE, italic=True)


def render_markdown_body(doc, markdown, *, skip_title=True):
    lines = markdown.splitlines()
    skipped_first_h1 = not skip_title
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# ") and not skipped_first_h1:
            skipped_first_h1 = True
            continue
        if line.startswith("**Phiên bản:") or line.startswith("**Cập nhật:") or line.startswith("**Bot") or line.startswith("**Website:") or line.startswith("**Định hướng:"):
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_run_font(p.add_run(line[2:]), size=10.5, color=MUTED, italic=True)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, line[3:])
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, line[4:])
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line)


def register_pdf_fonts():
    font_dir = Path("C:/Windows/Fonts")
    pdfmetrics.registerFont(TTFont("AASArial", str(font_dir / "arial.ttf")))
    pdfmetrics.registerFont(TTFont("AASArialBold", str(font_dir / "arialbd.ttf")))
    pdfmetrics.registerFont(TTFont("AASArialItalic", str(font_dir / "ariali.ttf")))
    pdfmetrics.registerFontFamily(
        "AASArial",
        normal="AASArial",
        bold="AASArialBold",
        italic="AASArialItalic",
        boldItalic="AASArialBold",
    )


def pdf_inline(text):
    parts = []
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        parts.append(escape(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            parts.append(f"<b>{escape(token[2:-2])}</b>")
        elif token.startswith("`"):
            parts.append(f"<font name='Courier'>{escape(token[1:-1])}</font>")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            parts.append(f"<link href='{escape(url)}' color='#08734F'><b>{escape(label)}</b></link>")
        cursor = match.end()
    parts.append(escape(text[cursor:]))
    return "".join(parts)


def pdf_styles(preset):
    styles = getSampleStyleSheet()
    accent = colors.HexColor("#08734F") if preset == "guide" else colors.HexColor("#2E74B5")
    dark = colors.HexColor("#053B2B") if preset == "guide" else colors.HexColor("#1F4D78")
    return {
        "title": ParagraphStyle(
            "AAS_Title",
            parent=styles["Title"],
            fontName="AASArialBold",
            fontSize=26 if preset == "guide" else 22,
            leading=31 if preset == "guide" else 27,
            textColor=dark,
            alignment=TA_LEFT,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "AAS_Subtitle",
            parent=styles["Normal"],
            fontName="AASArial",
            fontSize=12.5,
            leading=17,
            textColor=colors.HexColor("#5B6862"),
            spaceAfter=16,
        ),
        "meta": ParagraphStyle(
            "AAS_Meta",
            parent=styles["Normal"],
            fontName="AASArial",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#35463F"),
            leftIndent=8,
            spaceAfter=3,
        ),
        "h1": ParagraphStyle(
            "AAS_H1",
            parent=styles["Heading1"],
            fontName="AASArialBold",
            fontSize=15,
            leading=19,
            textColor=accent,
            spaceBefore=13,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "AAS_H2",
            parent=styles["Heading2"],
            fontName="AASArialBold",
            fontSize=12.5,
            leading=16,
            textColor=dark,
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "AAS_Body",
            parent=styles["BodyText"],
            fontName="AASArial",
            fontSize=10.2,
            leading=14.5 if preset == "guide" else 14,
            textColor=colors.HexColor("#18241F"),
            spaceAfter=5.5,
        ),
        "bullet": ParagraphStyle(
            "AAS_Bullet",
            parent=styles["BodyText"],
            fontName="AASArial",
            fontSize=10.2,
            leading=14.5,
            textColor=colors.HexColor("#18241F"),
            leftIndent=18,
            firstLineIndent=-9,
            bulletIndent=5,
            spaceAfter=3.5,
        ),
        "note": ParagraphStyle(
            "AAS_Note",
            parent=styles["BodyText"],
            fontName="AASArialItalic",
            fontSize=9.7,
            leading=14,
            textColor=dark,
            backColor=colors.HexColor("#E8EEF5" if preset == "terms" else "#E7F4EE"),
            borderPadding=9,
            borderColor=colors.HexColor("#C8D8D0"),
            borderWidth=0.5,
            spaceBefore=5,
            spaceAfter=10,
        ),
    }


def pdf_header_footer(canvas, doc, label):
    canvas.saveState()
    canvas.setFont("AASArial", 8.2)
    canvas.setFillColor(colors.HexColor("#5B6862"))
    canvas.drawString(doc.leftMargin, 0.55 * inch, label)
    canvas.drawRightString(letter[0] - doc.rightMargin, 0.55 * inch, f"Trang {doc.page}")
    canvas.restoreState()


def build_pdf_from_markdown(markdown_path, output_path, *, preset, title, subtitle, metadata):
    register_pdf_fonts()
    styles = pdf_styles(preset)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.82 * inch,
        rightMargin=0.82 * inch,
        topMargin=0.78 * inch,
        bottomMargin=0.78 * inch,
        title=title,
        author="TOAN AAS",
    )
    story = [
        Paragraph(pdf_inline(title), styles["title"]),
        Paragraph(pdf_inline(subtitle), styles["subtitle"]),
    ]
    for label, value in metadata:
        story.append(Paragraph(f"<b>{escape(label)}:</b> {escape(value)}", styles["meta"]))
    story.append(Spacer(1, 9))

    markdown = markdown_path.read_text(encoding="utf-8")
    first_h1_skipped = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# ") and not first_h1_skipped:
            first_h1_skipped = True
            continue
        if line.startswith("**Phiên bản:") or line.startswith("**Cập nhật:") or line.startswith("**Bot") or line.startswith("**Website:") or line.startswith("**Định hướng:"):
            continue
        if line.startswith("> "):
            story.append(Paragraph(pdf_inline(line[2:]), styles["note"]))
        elif line.startswith("## "):
            story.append(Paragraph(pdf_inline(line[3:]), styles["h1"]))
        elif line.startswith("### "):
            story.append(Paragraph(pdf_inline(line[4:]), styles["h2"]))
        elif line.startswith("- "):
            story.append(Paragraph(pdf_inline(line[2:]), styles["bullet"], bulletText="•"))
        elif re.match(r"^\d+\.\s", line):
            number = re.match(r"^(\d+)\.", line).group(1)
            text = re.sub(r"^\d+\.\s+", "", line)
            story.append(Paragraph(pdf_inline(text), styles["bullet"], bulletText=f"{number}."))
        else:
            story.append(Paragraph(pdf_inline(line), styles["body"]))

    label = "TOAN AAS | Hướng dẫn sử dụng V2" if preset == "guide" else "TOAN AAS | Điều khoản & Chính sách V2"
    doc.build(
        story,
        onFirstPage=lambda canvas, document: pdf_header_footer(canvas, document, label),
        onLaterPages=lambda canvas, document: pdf_header_footer(canvas, document, label),
    )


def build_guide():
    markdown = GUIDE_MD.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc, "TOAN AAS | Hướng dẫn sử dụng V2")
    configure_styles(doc, "guide")
    add_title_block(
        doc,
        title="TOAN AAS",
        subtitle="Hướng dẫn sử dụng cho khách hàng",
        metadata=[
            ("Phiên bản", "V2 - cập nhật 23/06/2026"),
            ("Bot Telegram", "@toanaasbot"),
            ("Website", "www.toanaas.vn"),
            ("Định hướng", "AI Automation System / Content Factory / Bot hỗ trợ công việc"),
        ],
        preset="guide",
    )
    render_markdown_body(doc, markdown)
    doc.core_properties.title = "TOAN AAS - Hướng dẫn sử dụng cho khách hàng V2"
    doc.core_properties.subject = "Hướng dẫn public TOAN AAS"
    doc.core_properties.author = "TOAN AAS"
    doc.save(GUIDE_DOCX)


def build_terms():
    markdown = TERMS_MD.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc, "TOAN AAS | Điều khoản & Chính sách V2")
    configure_styles(doc, "terms")
    add_title_block(
        doc,
        title="TOAN AAS",
        subtitle="Điều khoản & Chính sách dịch vụ",
        metadata=[
            ("Phiên bản", "V2"),
            ("Cập nhật", "13/06/2026"),
            ("Website", "www.toanaas.vn"),
            ("Hỗ trợ", "@toanaas | @toanaasbot"),
        ],
        preset="terms",
    )
    render_markdown_body(doc, markdown)
    doc.core_properties.title = "TOAN AAS - Điều khoản & Chính sách dịch vụ V2"
    doc.core_properties.subject = "Điều khoản và chính sách public TOAN AAS"
    doc.core_properties.author = "TOAN AAS"
    doc.save(TERMS_DOCX)


def main():
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    build_guide()
    build_terms()
    build_pdf_from_markdown(
        GUIDE_MD,
        GUIDE_PREVIEW_PDF,
        preset="guide",
        title="TOAN AAS",
        subtitle="Hướng dẫn sử dụng cho khách hàng",
        metadata=[
            ("Phiên bản", "V2 - cập nhật 23/06/2026"),
            ("Bot Telegram", "@toanaasbot"),
            ("Website", "www.toanaas.vn"),
            ("Định hướng", "AI Automation System / Content Factory / Bot hỗ trợ công việc"),
        ],
    )
    build_pdf_from_markdown(
        TERMS_MD,
        TERMS_PDF,
        preset="terms",
        title="TOAN AAS",
        subtitle="Điều khoản & Chính sách dịch vụ",
        metadata=[
            ("Phiên bản", "V2"),
            ("Cập nhật", "13/06/2026"),
            ("Website", "www.toanaas.vn"),
            ("Hỗ trợ", "@toanaas | @toanaasbot"),
        ],
    )
    print(GUIDE_DOCX)
    print(TERMS_DOCX)
    print(GUIDE_PREVIEW_PDF)
    print(TERMS_PDF)


if __name__ == "__main__":
    main()
